#!/usr/bin/env python3
"""Open every file a sweep produced and say whether it is usable.

    /usr/bin/python3 verify_artifacts.py [--dir DIR] [--out REPORT.md] [--shots DIR]

This is the half of round 15's axis that a model cannot do. The sweep model can
read a file and describe it; it cannot render one and see whether the chart drew.
That distinction is the whole reason the round exists: for a day every chart the
fleet wrote was a title, an empty bordered box and a console error, while the
tool reply gave a true path and a true byte count, and reading the file would
have shown a `<div class="plotly-graph-div">` sitting right there looking
correct. Only rendering it shows the div is empty.

So each artifact gets checked three ways, and each way can fail on its own:

  stands alone   every src= and href= in the page markup is walked. Script
                 BODIES are stripped before the search, because the minified
                 plotting library mentions its own website in its source and
                 matching that reads as an external dependency when it is a
                 comment. Then the file is copied ALONE into an empty directory,
                 which is where a sidecar dependency stops being invisible.

  renders        chromium opens the copy with every non-file:// request aborted,
                 so anything fetched from the network fails rather than quietly
                 succeeding on a machine that happens to be online. Plotly
                 figures are counted twice: containers present, and containers
                 that actually contain an <svg class="main-svg">. divs=1 drawn=0
                 is the exact signature of the regression this round follows.

  is its type    a .csv parses, a .xlsx opens with its sheets, a .docx has
                 paragraphs, a .pptx has slides, a .pdf has pages, an archive
                 lists its members. A file of the right size and the wrong
                 insides passes every check that looks at the reply.

Exit code is 1 if any artifact failed, so this can gate a round.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import tempfile
import traceback
from pathlib import Path

# Script bodies are removed before the reference search: plotly.min.js contains
# `https://plotly.com/` in its own source, and a naive scan reports every
# self-contained page as depending on the network.
#
# The opening tag is KEPT. Replacing the whole element deletes the `src` of
# `<script src="plotly.min.js"></script>` along with the body, which is the one
# reference this file exists to find -- the self-check caught exactly that, on
# the first version of this line.
_SCRIPT_BODY = re.compile(r"(<script\b[^>]*>).*?</script>", re.I | re.S)
_REF = re.compile(r"""\b(?:src|href)\s*=\s*["']([^"']+)["']""", re.I)
_INERT_PREFIXES = ("data:", "#", "javascript:", "about:", "mailto:")

# Files the servers write beside an artifact rather than as one. They are
# checked as their own JSON, not as a dependency of the file they sit next to.
_SIDECAR_SUFFIXES = (".mcp_receipt.json",)


def external_refs(html: str) -> list[str]:
    markup = _SCRIPT_BODY.sub(r"\1</script>", html)
    return [u for u in _REF.findall(markup) if not u.lower().startswith(_INERT_PREFIXES)]


def check_html(path: Path, shots: Path | None, browser) -> dict:
    """Copy the page out alone, render it with the network off, report what drew."""
    text = path.read_text(encoding="utf-8", errors="replace")
    result: dict = {
        "kind": "html",
        "refs": external_refs(text),
        "divs": text.count('class="plotly-graph-div"'),
    }

    with tempfile.TemporaryDirectory() as tmp:
        alone = Path(tmp) / path.name
        shutil.copy2(path, alone)

        blocked: list[str] = []
        errors: list[str] = []
        page = browser.new_page(viewport={"width": 1280, "height": 900})

        def route(r):
            # file:// is the page itself. Everything else is refused, so a page
            # that needs the network fails here instead of passing on a machine
            # that happens to have it.
            if r.request.url.startswith("file://"):
                r.continue_()
            else:
                blocked.append(r.request.url)
                r.abort()

        page.route("**/*", route)
        page.on("pageerror", lambda e: errors.append(str(e)))
        page.on(
            "console",
            lambda m: errors.append(f"console: {m.text}") if m.type == "error" else None,
        )
        try:
            page.goto(alone.as_uri(), wait_until="load", timeout=45_000)
            page.wait_for_timeout(3_000)
            result["drawn"] = page.locator(".plotly-graph-div svg.main-svg").count()
            result["divs"] = max(result["divs"], page.locator(".plotly-graph-div").count())
            result["text_len"] = len(page.locator("body").inner_text())
            if shots is not None:
                shot = shots / (path.stem + ".png")
                page.screenshot(path=str(shot))
                result["shot"] = str(shot)
        except Exception as exc:
            result["render_error"] = f"{type(exc).__name__}: {exc}"
        finally:
            page.close()

        result["blocked"] = sorted(set(blocked))
        result["errors"] = errors[:5]

    problems = []
    if result["refs"]:
        problems.append(f"needs {len(result['refs'])} file(s)/URL(s) outside itself: {result['refs'][:3]}")
    if result["blocked"]:
        problems.append(f"fetched {len(result['blocked'])} URL(s) when opened: {result['blocked'][:3]}")
    if result.get("render_error"):
        problems.append(result["render_error"])
    if result["divs"] and not result.get("drawn"):
        problems.append(f"{result['divs']} chart container(s), none drew")
    elif result["divs"] and result["drawn"] < result["divs"]:
        problems.append(f"{result['drawn']} of {result['divs']} charts drew")
    if not result["divs"] and result.get("text_len", 1) == 0:
        problems.append("page rendered empty")
    if result["errors"]:
        problems.append(f"js error: {result['errors'][0][:90]}")
    result["problems"] = problems
    return result


def check_tabular(path: Path) -> dict:
    import pandas as pd

    df = pd.read_csv(path)
    return {"kind": "csv", "rows": len(df), "cols": len(df.columns), "problems": []}


def check_xlsx(path: Path) -> dict:
    import openpyxl

    wb = openpyxl.load_workbook(path)
    sheets = wb.sheetnames
    rows = sum(ws.max_row for ws in wb.worksheets)
    return {"kind": "xlsx", "sheets": len(sheets), "rows": rows, "problems": []}


def check_docx(path: Path) -> dict:
    import docx

    d = docx.Document(str(path))
    paras = [p for p in d.paragraphs if p.text.strip()]
    return {"kind": "docx", "paragraphs": len(paras), "tables": len(d.tables), "problems": []}


def check_pptx(path: Path) -> dict:
    from pptx import Presentation

    p = Presentation(str(path))
    return {"kind": "pptx", "slides": len(p.slides), "problems": []}


def check_pdf(path: Path) -> dict:
    from pypdf import PdfReader

    r = PdfReader(str(path))
    return {"kind": "pdf", "pages": len(r.pages), "problems": []}


def check_archive(path: Path) -> dict:
    import tarfile
    import zipfile

    if path.suffix == ".zip":
        with zipfile.ZipFile(path) as z:
            bad = z.testzip()
            members = len(z.namelist())
        return {"kind": "zip", "members": members, "problems": [f"corrupt member {bad}"] if bad else []}
    with tarfile.open(path) as t:
        members = len(t.getnames())
    return {"kind": "tar", "members": members, "problems": []}


def check_json(path: Path) -> dict:
    data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    size = len(data) if isinstance(data, (list, dict)) else 1
    return {"kind": "json", "entries": size, "problems": []}


def check_text(path: Path) -> dict:
    text = path.read_text(encoding="utf-8", errors="replace")
    return {
        "kind": path.suffix.lstrip(".") or "text",
        "chars": len(text),
        "problems": [] if text.strip() else ["file is empty"],
    }


_HANDLERS = {
    ".csv": check_tabular,
    ".xlsx": check_xlsx,
    ".xlsm": check_xlsx,
    ".docx": check_docx,
    ".pptx": check_pptx,
    ".pdf": check_pdf,
    ".zip": check_archive,
    ".tar": check_archive,
    ".gz": check_archive,
    ".tgz": check_archive,
    ".json": check_json,
    ".md": check_text,
    ".txt": check_text,
}


def self_check() -> int:
    """Prove each way of failing still fires, on files built to fail that way.

    Written because the tests that were supposed to catch the sidecar
    regression passed all the way through it: they asserted the sidecar tag was
    present, so they could not fail on the thing that was wrong. A checker that
    has never been seen to fail is a checker nobody has tested, and this one is
    about to be used to sign off a round.
    """
    from playwright.sync_api import sync_playwright

    tmp = Path(tempfile.mkdtemp(prefix="artifact-selfcheck-"))
    cases: list[tuple[str, Path, str]] = []

    # 1. The regression itself: the library lives in a sibling file. Copied out
    #    alone the container is there and nothing draws.
    (tmp / "plotly.min.js").write_text("window.Plotly={newPlot:function(){}};", encoding="utf-8")
    sidecar = tmp / "sidecar_chart.html"
    sidecar.write_text(
        '<script src="plotly.min.js"></script>'
        '<div id="c" class="plotly-graph-div"></div>'
        '<script>Plotly.newPlot("c",[{"x":[1],"y":[2],"type":"bar"}],{});</script>',
        encoding="utf-8",
    )
    cases.append(("a page whose library is a sibling file", sidecar, "outside itself"))

    # 2. Fetched at open time -- the shape a geo map has.
    remote = tmp / "remote_chart.html"
    remote.write_text(
        '<div class="plotly-graph-div"></div>'
        '<script>fetch("https://example.com/world.json");</script>',
        encoding="utf-8",
    )
    cases.append(("a page that fetches when opened", remote, "fetched"))

    # 3. Right extension, wrong insides.
    fake = tmp / "not_really.xlsx"
    fake.write_bytes(b"this is not a workbook")
    cases.append(("a .xlsx that is not one", fake, "will not open"))

    # 4. Written, reported, empty.
    blank = tmp / "blank_report.md"
    blank.write_text("   \n", encoding="utf-8")
    cases.append(("a report with nothing in it", blank, "empty"))

    failures = 0
    with sync_playwright() as pw:
        browser = pw.chromium.launch(args=["--no-sandbox"])
        for name, path, expected in cases:
            if path.suffix.lower() in (".html", ".htm"):
                res = check_html(path, None, browser)
            else:
                try:
                    res = _HANDLERS[path.suffix.lower()](path)
                except Exception as exc:
                    res = {"problems": [f"will not open as {path.suffix}: {type(exc).__name__}: {exc}"]}
            joined = "; ".join(res["problems"])
            hit = expected in joined
            print(f"{'ok  ' if hit else 'MISS'}  {name}")
            print(f"        wanted {expected!r} in: {joined or '(no problems reported)'}")
            failures += 0 if hit else 1
        browser.close()

    # And one that must come back clean, or every result above is just a
    # checker that says FAIL to everything.
    good = tmp / "good.md"
    good.write_text("# real content\n", encoding="utf-8")
    clean = check_text(good)
    if clean["problems"]:
        print(f"MISS  a good file was called broken: {clean['problems']}")
        failures += 1
    else:
        print("ok    a good file still passes")

    shutil.rmtree(tmp, ignore_errors=True)
    print(f"\nself-check: {len(cases) + 1 - failures} of {len(cases) + 1} behaved")
    return 1 if failures else 0


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--dir", type=Path, default=Path("/root/Harnesses/data"))
    ap.add_argument("--out", type=Path)
    ap.add_argument("--shots", type=Path)
    ap.add_argument("--skip", default="Ad_Data.csv", help="comma-separated names to leave alone")
    ap.add_argument("--self-check", action="store_true", help="prove the checks can still fail")
    args = ap.parse_args()

    if args.self_check:
        return self_check()

    skip = {s.strip() for s in args.skip.split(",") if s.strip()}
    if args.shots:
        args.shots.mkdir(parents=True, exist_ok=True)

    files = [
        p
        for p in sorted(args.dir.rglob("*"))
        if p.is_file() and p.name not in skip and ".mcp_versions" not in p.parts
    ]
    if not files:
        print(f"no files under {args.dir}", file=sys.stderr)
        return 1

    from playwright.sync_api import sync_playwright

    rows: list[tuple[Path, dict]] = []
    with sync_playwright() as pw:
        browser = pw.chromium.launch(args=["--no-sandbox"])
        for p in files:
            suffix = p.suffix.lower()
            try:
                if suffix in (".html", ".htm"):
                    res = check_html(p, args.shots, browser)
                elif p.name.endswith(_SIDECAR_SUFFIXES):
                    res = check_json(p)
                elif suffix in _HANDLERS:
                    res = _HANDLERS[suffix](p)
                else:
                    res = {"kind": suffix.lstrip(".") or "?", "problems": [], "note": "not opened"}
            except Exception as exc:
                res = {
                    "kind": suffix.lstrip(".") or "?",
                    "problems": [f"will not open as {suffix or 'its type'}: {type(exc).__name__}: {exc}"],
                    "trace": traceback.format_exc(limit=2),
                }
            res["bytes"] = p.stat().st_size
            rows.append((p, res))
            mark = "FAIL" if res["problems"] else "ok"
            print(f"{mark:>4}  {p.relative_to(args.dir)}  ({res['kind']}, {res['bytes']:,} B)")
            for problem in res["problems"]:
                print(f"        {problem}")

    failed = [(p, r) for p, r in rows if r["problems"]]

    if args.out:
        lines = [
            f"# Artifact check — {len(rows)} files under {args.dir}",
            "",
            f"**{len(failed)} failed, {len(rows) - len(failed)} ok**",
            "",
            "| file | type | bytes | verdict | detail |",
            "|---|---|---|---|---|",
        ]
        for p, r in rows:
            detail = "; ".join(r["problems"]) if r["problems"] else _summary(r)
            verdict = "FAIL" if r["problems"] else "ok"
            lines.append(
                f"| `{p.relative_to(args.dir)}` | {r['kind']} | {r['bytes']:,} | {verdict} | {detail} |"
            )
        args.out.write_text("\n".join(lines) + "\n", encoding="utf-8")
        print(f"\nwrote {args.out}")

    print(f"\n{len(failed)} failed of {len(rows)}")
    return 1 if failed else 0


def _summary(r: dict) -> str:
    keep = ("rows", "cols", "sheets", "paragraphs", "tables", "slides", "pages", "members", "entries", "chars")
    bits = [f"{k}={r[k]}" for k in keep if k in r]
    if r.get("divs"):
        bits.append(f"charts {r.get('drawn', 0)}/{r['divs']} drew")
    return ", ".join(bits) or "opened"


if __name__ == "__main__":
    sys.exit(main())
